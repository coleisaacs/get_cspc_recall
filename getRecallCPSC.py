#!/usr/bin/env pyython
# -*- coding: utf-8 -*-

"""

__description__  = "Pulls recall information from the Consumer Product
                    Safety Commision (CSPC) and outputs to a formatted (.docx)."
__author__ = "Cole Isaacs"
__credits__ = "Cole Isaacs"
__license__ = "GNU GPLv3 except dependencies where their licenses apply."
__dependencies__ = "CSCPC API, Python packages: python-docx, requests, imageio."
__version__ = "1.0.1"
__maintainer__ = "Cole Isaacs"
__email__ = "coleisaacs@gmail.com"
__status__ = "Production"
__notes__  = "CSCPC API: https://cpsc.gov/Recalls/
                         CPSC-Recalls-Application-Program-Interface-API-Information
             
             Produced .docx requires some editing. 
             "

"""

# Import dependencies.

import json
import requests
import imageio
from datetime import date
from docx import Document
from docx.shared import Inches

# Get current date and rewind one year for starting timeframe of recall list.

getCurDate = date.today()
year = getCurDate.year - 1
month = getCurDate.month
day =  getCurDate.day

prepDate = str(year) + "-" + str(month) + "-" + str(day)

# CSCPC API URI parameters.
def make_call():
    """
    Make API call. 

    """
    global makeCall
    apiCSCPC = "http://www.saferproducts.gov/RestWebServices/Recall?"
    apiParam = "format=json&RecallDateStart=" + str(finDate)
    apiCall = apiCSCPC + apiParam
    
    makeCall = requests.get(apiCall).json()

    return makeCall


#  Filename to save to for API output
saveOut = 'recall_' + str(getCurDate) + '.docx'

def get_date(prepDate):
    """
    Check date and if not OK, prompt to correct.

    """
    global finDate     
    
    while True:
        print("Use " + prepDate + " as the date?")
        ans = input("Enter Y or y for Yes and N or n for No: ")


        if ('y' == ans) or \
           ('Y' == ans) or \
           ('n' == ans) or \
           ('N' == ans):

           # If valid no action.
           finDate = prepDate
           print("OK. ")
           
           break


    if ("n" == ans) or ("N" == ans):

    # Validate user entry.

            while True:
                year = input("Enter year (4-digit/YYYY): ")
                month = input("Enter month (2-digit/MM): ")
                day =  input ("Enter day (2-digit/DD): ")

                try:
                    year = int(year)
                    month = int(month)
                    day = int(day)


                except:
                    print(
                        "Please enter 4 digits for year, \
                         and 2 digits for month and day. "
                        )
                    continue

                if (0 >= month) or \
                   (13 <= month) or \
                   (0 >= day) or \
                   (32 <= day) or \
                   (getCurDate.year < year):

                    # If out of range restart loop

                    print("Invalid entry.")
                    continue

                else:

                     if(2 == len(str(month))) and (2 == len(str(day))):
                          # Set date to string and format for API call.
                          year = str(year)
                          month = str(month)
                          day = str(day)

                          finDate = year + "-" + month + "-" + day
                        
                     else:
                         # Set date to string and format for API call.
                         year = str(year)
                         month = str(month)
                         day = str(day)

                         finDate = year + "-0" + month + "-0" + day
                break

            return finDate

def create_docx():
    """
    Parse JSON response and output to a docx.

    """

    print("Creating .docx from: " + str(finDate) + ". Please wait...")

    # Get fresh call. 
    make_call()

    recs = len(makeCall)
    rn = 1

    docx = Document()
   

    

    print("Total Recalls: " + str(recs))
    print("")

    for i in makeCall:
        # Display that script is still processing.

        print("Processing: " + str(rn) + "/" + \
              str(recs) + "  ======== " + \
              str(i["Products"][0]["Name"]))

        # Create docx.

        p = docx.add_paragraph("")
        p.add_run(i["Products"][0]["Name"]).bold = True

        p = docx.add_paragraph("")
        p.add_run("Hazard: ").bold = True
        p.add_run(i["Hazards"][0]["Name"])

        p = docx.add_paragraph("")
        p = docx.add_paragraph(i["Description"])

        try:
            p = docx.add_paragraph(i["Retailers"][0]["Name"])
            
        except:
            p = docx.add_paragraph("")

        p = docx.add_paragraph(i["ConsumerContact"])

        p = docx.add_paragraph("")
        p.add_run("Remedy: ").bold = True

        try:
            p.add_run(i["Remedies"][0]["Name"])
            
        except:
            p = docx.add_paragraph("")


        # Iterate through image URLs and get content.

        for x in i["Images"]:
            
            # Check for errors with URL
            try:
                imge = imageio.imread(x["URL"])
                imageio.imwrite('imge.png', imge)
            
            except: 
                p = docx.add_paragraph(x["URL"])

            else:
                p = docx.add_picture('imge.png')
        
        rn += 1

    docx.save(saveOut)

print("Get Recalls from CPSC.gov")
print("========================")

get_date(prepDate)

create_docx()

print("========================")
